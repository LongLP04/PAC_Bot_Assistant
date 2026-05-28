from pathlib import Path

import pandas as pd
from docx import Document


SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".docx",
    ".xlsx",
    ".csv",
}


def is_supported_file(file_path: str) -> bool:
    """
    Kiểm tra file có thuộc định dạng được hỗ trợ hay không.
    """
    suffix = Path(file_path).suffix.lower()
    return suffix in SUPPORTED_EXTENSIONS


def read_txt_file(file_path: Path) -> str:
    """
    Đọc file .txt.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def read_md_file(file_path: Path) -> str:
    """
    Đọc file .md.
    Markdown cũng là text nên đọc giống .txt.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def read_docx_file(file_path: Path) -> str:
    """
    Đọc file .docx.
    Bao gồm nội dung đoạn văn và bảng trong Word.
    """
    document = Document(file_path)
    parts = []

    # Đọc đoạn văn
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Đọc bảng trong Word
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n[BẢNG WORD {table_index}]")

        for row in table.rows:
            row_values = []

            for cell in row.cells:
                value = cell.text.strip().replace("\n", " ")
                row_values.append(value)

            if any(row_values):
                parts.append(" | ".join(row_values))

    return "\n".join(parts)


def read_csv_file(file_path: Path) -> str:
    """
    Đọc file .csv.
    Thử nhiều kiểu encoding để hạn chế lỗi tiếng Việt.
    """
    try:
        df = pd.read_csv(file_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(file_path, encoding="utf-8-sig")
    except Exception:
        df = pd.read_csv(file_path, encoding="cp1258")

    return dataframe_to_text(df, title="BẢNG CSV")


def read_xlsx_file(file_path: Path) -> str:
    """
    Đọc file .xlsx.
    Đọc toàn bộ sheet trong file Excel.
    """
    excel_file = pd.ExcelFile(file_path)
    parts = []

    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        parts.append(
            dataframe_to_text(
                df,
                title=f"BẢNG EXCEL - SHEET: {sheet_name}"
            )
        )

    return "\n\n".join(parts)


def dataframe_to_text(df: pd.DataFrame, title: str = "BẢNG DỮ LIỆU") -> str:
    """
    Chuyển bảng Excel/CSV thành text để bot nạp vào knowledge.
    """
    if df.empty:
        return f"[{title}]\nKhông có dữ liệu."

    df = df.fillna("")

    max_rows = 300

    if len(df) > max_rows:
        df = df.head(max_rows)
        note = (
            f"\n[GHI CHÚ] File gốc có nhiều hơn {max_rows} dòng, "
            f"bot chỉ nạp {max_rows} dòng đầu tiên."
        )
    else:
        note = ""

    parts = [f"[{title}]{note}"]

    # Header
    columns = [str(col).strip() for col in df.columns]
    parts.append(" | ".join(columns))

    # Rows
    for _, row in df.iterrows():
        values = [str(row[col]).strip() for col in df.columns]
        parts.append(" | ".join(values))

    return "\n".join(parts)


def read_document(file_path: str) -> str:
    """
    Đọc tài liệu theo định dạng file.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return read_txt_file(path)

    if suffix == ".md":
        return read_md_file(path)

    if suffix == ".docx":
        return read_docx_file(path)

    if suffix == ".xlsx":
        return read_xlsx_file(path)

    if suffix == ".csv":
        return read_csv_file(path)

    raise ValueError(f"Định dạng file nạp vào không được hỗ trợ: {suffix}")